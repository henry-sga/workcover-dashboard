"""
Document generator — produces Register of Injury, Toolbox Talks, and other forms as DOCX files.
"""

from __future__ import annotations

import io
from datetime import date
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH


# --- Toolbox Talk content by injury type ---

TOOLBOX_TALK_CONTENT: dict[str, dict] = {
    "Manual Handling / Back": {
        "topic": "Safe Manual Handling & Back Injury Prevention",
        "hazards": [
            "Lifting heavy or awkward loads without assistance",
            "Repetitive bending, twisting or reaching movements",
            "Pushing/pulling loads across uneven surfaces",
            "Working in sustained awkward postures",
            "Fatigue leading to poor lifting technique",
        ],
        "controls": [
            "Use mechanical aids (trolleys, hoists, forklifts) wherever possible",
            "Follow the safe lifting technique: bend knees, keep load close, straight back, no twisting",
            "Team lift for any load over 16kg or awkward shape — ask for help",
            "Rotate tasks to avoid prolonged repetitive movements",
            "Store frequently used items between knee and shoulder height",
            "Report any discomfort or pain early — do not push through it",
        ],
        "discussion_points": [
            "What manual handling tasks do you do most often?",
            "Where can we use mechanical aids to reduce lifting?",
            "Do you know how to request help or report early discomfort?",
        ],
    },
    "Crush / Fracture": {
        "topic": "Crush & Fracture Injury Prevention",
        "hazards": [
            "Moving machinery or vehicles in the work area",
            "Unsecured loads that may shift or fall",
            "Working near heavy equipment without exclusion zones",
            "Inadequate guarding on machinery",
            "Poor housekeeping creating trip/fall hazards near hard surfaces",
        ],
        "controls": [
            "Maintain exclusion zones around moving plant and equipment",
            "Never place hands or body parts near pinch points or moving parts",
            "Ensure all machine guards are in place before operation",
            "Use lockout/tagout (LOTO) procedures during maintenance",
            "Wear appropriate PPE — steel cap boots, gloves, hard hat as required",
            "Secure all loads before transport and verify before releasing",
        ],
        "discussion_points": [
            "Are all machine guards in place and functional on your equipment?",
            "Do you know the lockout/tagout procedure for your area?",
            "Have you identified all pinch points in your workspace?",
        ],
    },
    "Laceration / Cut": {
        "topic": "Laceration & Cut Prevention — Sharp Object Safety",
        "hazards": [
            "Using blunt or damaged cutting tools",
            "Cutting towards the body or unsupported hand",
            "Broken glass, metal edges, or exposed sharp surfaces",
            "Inadequate PPE when handling sharp materials",
            "Rushing tasks involving knives, box cutters, or blades",
        ],
        "controls": [
            "Always cut away from your body and keep hands behind the blade",
            "Use safety knives with retractable blades where possible",
            "Replace blunt blades — they require more force and slip more easily",
            "Wear cut-resistant gloves when handling sharp materials",
            "Dispose of sharps safely in designated containers",
            "Keep work areas clean and free of broken glass or metal offcuts",
        ],
        "discussion_points": [
            "Are you using the correct cutting tools for the job?",
            "Do you have access to cut-resistant gloves?",
            "Where are the sharps disposal containers in your area?",
        ],
    },
    "Sprain / Strain": {
        "topic": "Sprain & Strain Prevention — Musculoskeletal Safety",
        "hazards": [
            "Overexertion from heavy or repetitive tasks",
            "Slippery or uneven walking surfaces",
            "Working in awkward or sustained postures",
            "Inadequate warm-up before physical tasks",
            "Fatigue reducing coordination and body awareness",
        ],
        "controls": [
            "Warm up and stretch before starting physical work",
            "Use correct techniques for lifting, pushing, and pulling",
            "Take regular micro-breaks during repetitive tasks",
            "Wear appropriate footwear with good grip and ankle support",
            "Report wet, uneven, or damaged floor surfaces immediately",
            "Rotate between different tasks throughout the shift",
        ],
        "discussion_points": [
            "Are there tasks that cause you fatigue or discomfort?",
            "Do you take regular breaks during repetitive work?",
            "Are walking surfaces in your area well maintained?",
        ],
    },
    "Chemical": {
        "topic": "Chemical Safety & Hazardous Substance Handling",
        "hazards": [
            "Skin or eye contact with hazardous chemicals",
            "Inhalation of fumes, vapours, or dust",
            "Mixing incompatible chemicals",
            "Incorrect storage or labelling of chemicals",
            "Spills without proper containment or cleanup procedures",
        ],
        "controls": [
            "Read the Safety Data Sheet (SDS) before using any chemical",
            "Wear required PPE — gloves, goggles, respirator as specified on SDS",
            "Use chemicals only in well-ventilated areas",
            "Never mix chemicals unless specifically instructed",
            "Label all containers including decanted chemicals",
            "Know the location of spill kits, eyewash stations, and emergency showers",
        ],
        "discussion_points": [
            "Do you know where the SDS register is for your area?",
            "Are all chemicals properly labelled and stored?",
            "Do you know how to use the spill kit and eyewash station?",
        ],
    },
    "Burns": {
        "topic": "Burn Prevention — Thermal & Chemical Burns",
        "hazards": [
            "Contact with hot surfaces, liquids, or steam",
            "Splashing of hot or corrosive liquids",
            "Exposure to open flames or radiant heat",
            "Chemical burns from acids, alkalis, or solvents",
            "Electrical equipment generating heat",
        ],
        "controls": [
            "Use heat-resistant gloves and PPE when working near hot surfaces",
            "Allow hot equipment and liquids to cool before handling",
            "Keep flammable materials away from heat sources",
            "Use splash guards and face shields when handling hot/corrosive liquids",
            "Know the location of first aid burns kits and emergency showers",
            "Follow correct procedures for heating, cooking, or welding operations",
        ],
        "discussion_points": [
            "Are hot surfaces in your area clearly marked?",
            "Do you know first aid treatment for burns?",
            "Are there any tasks where you feel at risk of burns?",
        ],
    },
    "Slip / Trip / Fall": {
        "topic": "Slip, Trip & Fall Prevention",
        "hazards": [
            "Wet, oily, or contaminated floor surfaces",
            "Loose cables, hoses, or clutter in walkways",
            "Uneven surfaces, damaged flooring, or missing handrails",
            "Poor lighting in walkways and stairwells",
            "Working at heights without fall protection",
        ],
        "controls": [
            "Clean up spills immediately and use wet floor signs",
            "Keep walkways clear of clutter, cables, and obstructions",
            "Report damaged flooring, loose tiles, or missing handrails",
            "Wear footwear with slip-resistant soles",
            "Use handrails on stairs and maintain 3 points of contact on ladders",
            "Ensure adequate lighting in all work and walkway areas",
        ],
        "discussion_points": [
            "Are there areas in your workplace with frequent spills or wet floors?",
            "Are walkways kept clear on every shift?",
            "Do you report damaged flooring or lighting issues?",
        ],
    },
    "Psychological": {
        "topic": "Psychological Safety & Mental Health in the Workplace",
        "hazards": [
            "Excessive workload or unrealistic deadlines",
            "Bullying, harassment, or conflict in the workplace",
            "Exposure to traumatic events or high-stress situations",
            "Lack of support, recognition, or control over work",
            "Social isolation or poor team communication",
        ],
        "controls": [
            "Speak up early if workload is unmanageable — talk to your supervisor",
            "Treat all colleagues with respect — zero tolerance for bullying",
            "Use the Employee Assistance Program (EAP) for confidential support",
            "Take regular breaks and use your leave entitlements",
            "Check in on your mates — ask R U OK?",
            "Report any concerns about workplace behaviour to your manager or HR",
        ],
        "discussion_points": [
            "Do you feel comfortable raising concerns with your supervisor?",
            "Do you know how to access the EAP?",
            "How can we better support each other as a team?",
        ],
    },
    "Disease / Illness": {
        "topic": "Workplace Health — Illness & Disease Prevention",
        "hazards": [
            "Exposure to biological hazards (mould, bacteria, viruses)",
            "Poor hygiene practices in shared facilities",
            "Prolonged exposure to dust, fumes, or airborne contaminants",
            "Inadequate ventilation in enclosed work areas",
            "Working while unwell and spreading illness",
        ],
        "controls": [
            "Wash hands regularly, especially before eating and after using facilities",
            "Stay home if you are unwell — do not come to work sick",
            "Use respiratory protection (masks) when working in dusty or contaminated areas",
            "Ensure ventilation systems are working and maintained",
            "Keep shared facilities and break rooms clean",
            "Report any unusual smells, mould, or contamination immediately",
        ],
        "discussion_points": [
            "Are shared facilities cleaned regularly?",
            "Do you know when respiratory protection is required?",
            "Are there any air quality concerns in your work area?",
        ],
    },
}

# Default content for injury types not specifically listed
_DEFAULT_TOOLBOX = {
    "topic": "General Workplace Safety",
    "hazards": [
        "Unfamiliar or new tasks without proper training",
        "Complacency and shortcuts in routine work",
        "Failure to report hazards or near-misses",
        "Not wearing required PPE",
        "Poor housekeeping and cluttered work areas",
    ],
    "controls": [
        "Follow all safe work procedures and site rules",
        "Wear required PPE at all times",
        "Report all hazards, near-misses, and incidents immediately",
        "Ask for help or clarification if you are unsure about a task",
        "Keep your work area clean and organised",
        "Take regular breaks and stay hydrated",
    ],
    "discussion_points": [
        "Are there any hazards in your area that haven't been reported?",
        "Do you feel confident in the safe work procedures for your tasks?",
        "Is there anything we can improve to make this workplace safer?",
    ],
}


def _add_row(table, label: str, value: str):
    """Add a label-value row to a table."""
    row = table.add_row()
    c0, c1 = row.cells
    c0.text = label
    c1.text = value or ""
    for cell in (c0, c1):
        for para in cell.paragraphs:
            para.style.font.size = Pt(10)


def generate_register_of_injury(incident_data: dict) -> bytes:
    """
    Generate a Register of Injury document (DOCX) and return the bytes.

    ``incident_data`` should contain keys like:
        worker_name, dob, email, phone, entity, site, state,
        date_of_injury, injury_description, body_part, nature_of_injury,
        treatment, witnesses, shift_structure, employment_type, tenure
    """
    doc = Document()

    # Title
    title = doc.add_heading("Register of Injury", level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph(f"Date Generated: {date.today().strftime('%d/%m/%Y')}")

    # --- Part A: Employee / Injured Person Details ---
    doc.add_heading("Part A — Employee / Injured Person Details", level=2)
    table_a = doc.add_table(rows=0, cols=2)
    table_a.style = "Table Grid"
    table_a.alignment = WD_TABLE_ALIGNMENT.CENTER

    _add_row(table_a, "Employee Name", incident_data.get("worker_name", ""))
    _add_row(table_a, "Date of Birth", incident_data.get("dob", ""))
    _add_row(table_a, "Employee Email", incident_data.get("email", ""))
    _add_row(table_a, "Employee Phone", incident_data.get("phone", ""))
    _add_row(table_a, "Employer / Entity", incident_data.get("entity", ""))
    _add_row(table_a, "Workplace / Site", incident_data.get("site", ""))
    _add_row(table_a, "State", incident_data.get("state", ""))
    _add_row(table_a, "Employment Type", incident_data.get("employment_type", ""))
    _add_row(table_a, "Tenure / Length of Service", incident_data.get("tenure", ""))
    _add_row(table_a, "Shift / Hours of Work", incident_data.get("shift_structure", ""))

    # --- Part B: Incident Details ---
    doc.add_heading("Part B — Incident Details", level=2)
    table_b = doc.add_table(rows=0, cols=2)
    table_b.style = "Table Grid"
    table_b.alignment = WD_TABLE_ALIGNMENT.CENTER

    _add_row(table_b, "Date of Injury / Incident", incident_data.get("date_of_injury", ""))
    _add_row(table_b, "Location of Incident", incident_data.get("site", ""))
    _add_row(table_b, "Description (What Happened)", incident_data.get("injury_description", ""))
    _add_row(table_b, "Nature of Injury", incident_data.get("nature_of_injury", ""))
    _add_row(table_b, "Body Part Injured", incident_data.get("body_part", ""))
    _add_row(table_b, "Treatment Given", incident_data.get("treatment", ""))
    _add_row(table_b, "Witness(es)", incident_data.get("witnesses", ""))

    # --- Part C: Declaration ---
    doc.add_heading("Part C — Declaration", level=2)
    doc.add_paragraph(
        "I declare that the information provided above is true and correct "
        "to the best of my knowledge."
    )

    table_c = doc.add_table(rows=0, cols=2)
    table_c.style = "Table Grid"
    table_c.alignment = WD_TABLE_ALIGNMENT.CENTER
    _add_row(table_c, "Employee Signature", "")
    _add_row(table_c, "Date", "")
    _add_row(table_c, "Manager / Supervisor Name", "")
    _add_row(table_c, "Manager Signature", "")
    _add_row(table_c, "Date", "")

    # Write to bytes
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.getvalue()


def generate_toolbox_talk(case_data: dict) -> bytes:
    """
    Generate a Toolbox Talk document (DOCX) pre-populated based on the injury type.

    ``case_data`` should contain keys like:
        worker_name, site, entity, state, date_of_injury,
        injury_description, injury_type
    """
    injury_type = case_data.get("injury_type", "")
    content = TOOLBOX_TALK_CONTENT.get(injury_type, _DEFAULT_TOOLBOX)

    doc = Document()

    # --- Header ---
    title = doc.add_heading("Toolbox Talk", level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run(content["topic"])
    run.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0x1E, 0x3A, 0x5F)

    # --- Details table ---
    details = doc.add_table(rows=0, cols=4)
    details.style = "Table Grid"
    details.alignment = WD_TABLE_ALIGNMENT.CENTER

    row = details.add_row()
    row.cells[0].text = "Date:"
    row.cells[1].text = date.today().strftime("%d/%m/%Y")
    row.cells[2].text = "Site:"
    row.cells[3].text = case_data.get("site", "") or ""

    row = details.add_row()
    row.cells[0].text = "Entity:"
    row.cells[1].text = case_data.get("entity", "") or ""
    row.cells[2].text = "Presented By:"
    row.cells[3].text = ""

    for r in details.rows:
        for cell in r.cells:
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.size = Pt(9)

    doc.add_paragraph("")

    # --- What happened (incident summary) ---
    doc.add_heading("Incident Summary", level=2)
    p = doc.add_paragraph()
    p.add_run("A workplace incident has occurred at this site. ").font.size = Pt(10)
    desc = case_data.get("injury_description", "")
    if desc:
        p.add_run(f"Details: {desc}").font.size = Pt(10)
    else:
        p.add_run(
            "The details are being reviewed. This toolbox talk addresses "
            "the key hazards and controls relevant to this type of incident."
        ).font.size = Pt(10)

    # --- Hazards ---
    doc.add_heading("Hazards Identified", level=2)
    for h in content["hazards"]:
        p = doc.add_paragraph(h, style="List Bullet")
        for run in p.runs:
            run.font.size = Pt(10)

    # --- Controls ---
    doc.add_heading("Controls & Safe Work Practices", level=2)
    for c in content["controls"]:
        p = doc.add_paragraph(c, style="List Bullet")
        for run in p.runs:
            run.font.size = Pt(10)

    # --- Discussion points ---
    doc.add_heading("Discussion Points", level=2)
    for d in content["discussion_points"]:
        p = doc.add_paragraph(d, style="List Bullet")
        for run in p.runs:
            run.font.size = Pt(10)

    # --- Attendance register ---
    doc.add_heading("Attendance Register", level=2)
    att = doc.add_table(rows=11, cols=3)
    att.style = "Table Grid"
    att.alignment = WD_TABLE_ALIGNMENT.CENTER
    att.rows[0].cells[0].text = "Name"
    att.rows[0].cells[1].text = "Signature"
    att.rows[0].cells[2].text = "Date"
    for cell in att.rows[0].cells:
        for para in cell.paragraphs:
            for run in para.runs:
                run.bold = True
                run.font.size = Pt(9)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.getvalue()


def generate_rtw_plan(case_data: dict, coc_data: dict | None = None) -> bytes:
    """Generate a Return to Work Plan document (DOCX)."""
    doc = Document()

    title = doc.add_heading("Return to Work Plan", level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph(f"Date Prepared: {date.today().strftime('%d/%m/%Y')}")

    doc.add_heading("1. Worker Details", level=2)
    t1 = doc.add_table(rows=0, cols=2)
    t1.style = "Table Grid"
    t1.alignment = WD_TABLE_ALIGNMENT.CENTER
    _add_row(t1, "Worker Name", case_data.get("worker_name", ""))
    _add_row(t1, "Employer / Entity", case_data.get("entity", ""))
    _add_row(t1, "Workplace / Site", case_data.get("site", ""))
    _add_row(t1, "State", case_data.get("state", ""))
    _add_row(t1, "Date of Injury", case_data.get("date_of_injury", ""))
    _add_row(t1, "Claim Number", case_data.get("claim_number", ""))

    doc.add_heading("2. Injury Details", level=2)
    t2 = doc.add_table(rows=0, cols=2)
    t2.style = "Table Grid"
    t2.alignment = WD_TABLE_ALIGNMENT.CENTER
    _add_row(t2, "Injury Type", case_data.get("injury_type", ""))
    _add_row(t2, "Injury Description", case_data.get("injury_description", ""))
    _add_row(t2, "Current Capacity", case_data.get("current_capacity", ""))

    doc.add_heading("3. Medical Restrictions (from COC)", level=2)
    t3 = doc.add_table(rows=0, cols=2)
    t3.style = "Table Grid"
    t3.alignment = WD_TABLE_ALIGNMENT.CENTER
    if coc_data:
        _add_row(t3, "Certificate Period", f"{coc_data.get('cert_from', '')} to {coc_data.get('cert_to', '')}")
        _add_row(t3, "Capacity", coc_data.get("capacity", ""))
        _add_row(t3, "Days per Week", str(coc_data.get("days_per_week", "") or ""))
        _add_row(t3, "Hours per Day", str(coc_data.get("hours_per_day", "") or ""))
    else:
        _add_row(t3, "Certificate Period", "No current COC on file")
        _add_row(t3, "Capacity", case_data.get("current_capacity", ""))
    _add_row(t3, "Restrictions / Limitations", "")
    _add_row(t3, "Activities to Avoid", "")

    doc.add_heading("4. Return to Work Duties", level=2)
    doc.add_paragraph(
        "The following modified / suitable duties have been identified in consultation "
        "with the worker, treating practitioner, and employer:"
    )
    t4 = doc.add_table(rows=6, cols=3)
    t4.style = "Table Grid"
    t4.alignment = WD_TABLE_ALIGNMENT.CENTER
    t4.rows[0].cells[0].text = "Task / Duty"
    t4.rows[0].cells[1].text = "Hours / Frequency"
    t4.rows[0].cells[2].text = "Restrictions"
    for cell in t4.rows[0].cells:
        for p in cell.paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(9)

    doc.add_heading("5. Work Schedule", level=2)
    t5 = doc.add_table(rows=2, cols=6)
    t5.style = "Table Grid"
    t5.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, day in enumerate(["", "Mon", "Tue", "Wed", "Thu", "Fri"]):
        t5.rows[0].cells[i].text = day
    t5.rows[1].cells[0].text = "Hours"
    for cell in t5.rows[0].cells:
        for p in cell.paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(9)

    doc.add_heading("6. Review & Goals", level=2)
    t6 = doc.add_table(rows=0, cols=2)
    t6.style = "Table Grid"
    t6.alignment = WD_TABLE_ALIGNMENT.CENTER
    _add_row(t6, "Plan Start Date", date.today().strftime("%d/%m/%Y"))
    _add_row(t6, "Next Review Date", "")
    _add_row(t6, "Short-term Goal (2 weeks)", "")
    _add_row(t6, "Medium-term Goal (6 weeks)", "")
    _add_row(t6, "Long-term Goal", "Return to full pre-injury duties")

    doc.add_heading("7. Agreement", level=2)
    doc.add_paragraph(
        "By signing below, all parties agree to this Return to Work Plan "
        "and commit to supporting the worker's safe and sustainable return to work."
    )
    t7 = doc.add_table(rows=0, cols=3)
    t7.style = "Table Grid"
    t7.alignment = WD_TABLE_ALIGNMENT.CENTER
    for role in ["Worker", "Employer / Manager", "Treating Practitioner", "Insurer / Rehab Provider"]:
        row = t7.add_row()
        row.cells[0].text = role
        row.cells[1].text = "Signature:"
        row.cells[2].text = "Date:"

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.getvalue()
