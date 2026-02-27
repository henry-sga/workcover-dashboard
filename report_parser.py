"""
Incident report parser — extracts text from uploaded PDF/DOCX files
and matches field labels to pre-fill the New Case wizard.
"""

from __future__ import annotations

import re
import io
from datetime import datetime
from typing import Optional


# ---------------------------------------------------------------------------
# Text extraction
# ---------------------------------------------------------------------------

def _extract_text_from_docx(file_bytes: bytes) -> str:
    """Extract all text from a .docx file (paragraphs + table cells)."""
    from docx import Document

    doc = Document(io.BytesIO(file_bytes))
    parts = []

    for para in doc.paragraphs:
        if para.text.strip():
            parts.append(para.text.strip())

    for table in doc.tables:
        for row in table.rows:
            row_text = []
            for cell in row.cells:
                cell_text = cell.text.strip()
                if cell_text:
                    row_text.append(cell_text)
            if row_text:
                parts.append(" | ".join(row_text))

    return "\n".join(parts)


def _extract_text_from_pdf(file_bytes: bytes) -> str:
    """Extract all text from a PDF using pdfplumber."""
    import pdfplumber

    text_parts = []
    with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
        for page in pdf.pages:
            page_text = page.extract_text()
            if page_text:
                text_parts.append(page_text)
    return "\n".join(text_parts)


# ---------------------------------------------------------------------------
# Field parsing helpers
# ---------------------------------------------------------------------------

def _find(pattern: str, text: str, flags=re.IGNORECASE) -> str | None:
    """Return the first captured group, or None."""
    m = re.search(pattern, text, flags)
    return m.group(1).strip() if m else None


def _parse_date(raw: str | None) -> str | None:
    """Try to normalise a date string to YYYY-MM-DD."""
    if not raw:
        return None
    for fmt in ("%d/%m/%Y", "%d-%m-%Y", "%d %B %Y", "%d %b %Y",
                "%Y-%m-%d", "%d/%m/%y", "%d-%m-%y"):
        try:
            return datetime.strptime(raw.strip(), fmt).strftime("%Y-%m-%d")
        except ValueError:
            continue
    return raw.strip()


# ---------------------------------------------------------------------------
# Main field extraction
# ---------------------------------------------------------------------------

def _parse_fields_from_text(text: str) -> dict:
    """
    Match common field labels from incident / Register of Injury reports.
    Returns a dict whose keys match the New Case wizard field names.

    Supports Australian SafeWork templates, employer templates, and common
    incident report formats.
    """
    fields: dict = {}

    # Employee name — multiple patterns for different form layouts
    name = (_find(r"(?:employee|worker|injured person|injured worker)(?:'s)?\s*(?:full\s*)?name\s*[:\-|]\s*(.+)", text)
            or _find(r"name\s+of\s+(?:employee|worker|injured person)\s*[:\-|]\s*(.+)", text)
            or _find(r"(?:surname|family\s+name)\s*[:\-|]\s*(.+)", text)
            or _find(r"name\s*[:\-|]\s*([A-Z][a-z]+(?:\s+[A-Z][a-z]+)+)", text))
    if name:
        fields["worker_name"] = name.split("\n")[0].strip()

    # Date of birth
    dob = _find(r"(?:date\s+of\s+birth|dob|d\.o\.b)\s*[:\-|]\s*(\S+)", text)
    if dob:
        fields["dob"] = _parse_date(dob) or dob

    # Phone
    phone = (_find(r"(?:phone|telephone|contact\s*number|mobile)\s*[:\-|]\s*([\d\s\+\(\)]{8,})", text)
             or _find(r"(?:ph|mob)\s*[:\-|]\s*([\d\s\+\(\)]{8,})", text))
    if phone:
        fields["phone"] = re.sub(r"\s+", " ", phone).strip()

    # Email
    email = (_find(r"(?:email|e-mail)\s*[:\-|]\s*([\w.\-+]+@[\w.\-]+\.\w+)", text)
             or _find(r"([\w.\-+]+@[\w.\-]+\.\w+)", text))
    if email:
        fields["email"] = email.strip()

    # Site / Workplace
    site = (_find(r"(?:workplace|work\s*site|site|location\s+of\s+workplace)\s*[:\-|]\s*(.+)", text)
            or _find(r"(?:place\s+of\s+incident|incident\s+location|location\s+of\s+incident)\s*[:\-|]\s*(.+)", text)
            or _find(r"(?:where\s+did\s+(?:the\s+)?(?:incident|injury|accident)\s+occur)\s*[:\-|?]\s*(.+)", text))
    if site:
        fields["site"] = site.split("\n")[0].strip()

    # Date of incident / injury — expanded patterns
    doi = (_find(r"(?:date\s+of\s+(?:injury|incident|accident|occurrence))\s*[:\-|]\s*(\S+)", text)
           or _find(r"(?:incident\s+date|injury\s+date|date\s+(?:incident|injury)\s+occurred)\s*[:\-|]\s*(\S+)", text)
           or _find(r"(?:when\s+did\s+(?:the\s+)?(?:incident|injury)\s+occur)\s*[:\-|?]\s*(\S+)", text)
           or _find(r"(?:date)\s*[:\-|]\s*(\d{1,2}[/\-]\d{1,2}[/\-]\d{2,4})", text))
    if doi:
        fields["date_of_injury"] = _parse_date(doi) or doi

    # Time of incident
    time_of = _find(r"(?:time\s+of\s+(?:incident|injury|accident))\s*[:\-|]\s*(\d{1,2}[:\.]?\d{0,2}\s*(?:am|pm)?)", text)
    if time_of:
        fields["time_of_incident"] = time_of.strip()

    # What happened / description — more flexible patterns
    desc = (_find(r"(?:what\s+happened|description\s+of\s+(?:injury|incident|accident)|how\s+(?:did\s+)?(?:the\s+)?(?:injury|incident)\s+occur|details?\s+of\s+(?:incident|injury|accident))\s*[:\-|]\s*(.+?)(?:\n\n|\n[A-Z])", text, re.IGNORECASE | re.DOTALL)
            or _find(r"(?:what\s+happened|description\s+of\s+(?:injury|incident)|brief\s+description|describe\s+(?:the\s+)?(?:incident|injury))\s*[:\-|?]\s*(.+)", text)
            or _find(r"(?:incident\s+details?|injury\s+details?)\s*[:\-|]\s*(.+?)(?:\n\n|\n[A-Z])", text, re.IGNORECASE | re.DOTALL))
    if desc:
        # Clean up: limit to reasonable length, trim trailing whitespace
        cleaned = desc.strip()
        if len(cleaned) > 500:
            cleaned = cleaned[:500] + "..."
        fields["injury_description"] = cleaned

    # Witnesses
    witness = (_find(r"(?:witness(?:es)?|witness\s+name)\s*[:\-|]\s*(.+)", text)
               or _find(r"(?:any\s+witnesses?)\s*[:\-|?]\s*(.+)", text))
    if witness:
        fields["witnesses"] = witness.split("\n")[0].strip()

    # Employment type
    emp_type = _find(r"(?:employment\s+(?:type|status|basis)|type\s+of\s+employment)\s*[:\-|]\s*(.+)", text)
    if emp_type:
        fields["employment_type"] = emp_type.split("\n")[0].strip()

    # Tenure / length of service
    tenure = (_find(r"(?:tenure|length\s+of\s+(?:service|employment))\s*[:\-|]\s*(.+)", text)
              or _find(r"(?:start\s+date|commencement\s+date|date\s+(?:of\s+)?commencement)\s*[:\-|]\s*(.+)", text))
    if tenure:
        fields["tenure"] = tenure.split("\n")[0].strip()

    # Hours / shift
    shift = (_find(r"(?:shift|hours\s+(?:of\s+work|worked)|roster)\s*[:\-|]\s*(.+)", text)
             or _find(r"(?:normal\s+working\s+hours|usual\s+hours|hours\s+per\s+week)\s*[:\-|]\s*(.+)", text))
    if shift:
        fields["shift_structure"] = shift.split("\n")[0].strip()

    # Nature of injury / body part
    nature = (_find(r"(?:nature\s+of\s+injury|type\s+of\s+injury|injury\s+type)\s*[:\-|]\s*(.+)", text)
              or _find(r"(?:mechanism\s+of\s+injury)\s*[:\-|]\s*(.+)", text))
    if nature:
        fields["nature_of_injury"] = nature.split("\n")[0].strip()

    body_part = (_find(r"(?:body\s+part|part\s+of\s+body|injured\s+body\s+part|area\s+affected)\s*[:\-|]\s*(.+)", text)
                 or _find(r"(?:area\s+of\s+injury|location\s+of\s+injury)\s*[:\-|]\s*(.+)", text))
    if body_part:
        fields["body_part"] = body_part.split("\n")[0].strip()

    # Treatment
    treatment = (_find(r"(?:treatment|medical\s+treatment|first\s+aid|treatment\s+given)\s*[:\-|]\s*(.+)", text)
                 or _find(r"(?:initial\s+treatment|action\s+taken)\s*[:\-|]\s*(.+)", text))
    if treatment:
        fields["treatment"] = treatment.split("\n")[0].strip()

    # Entity / employer — expanded
    entity = (_find(r"(?:employer|entity|company|business\s+name|organisation)\s*[:\-|]\s*(.+)", text)
              or _find(r"(?:name\s+of\s+employer|employer\s+name|trading\s+name)\s*[:\-|]\s*(.+)", text))
    if entity:
        fields["entity"] = entity.split("\n")[0].strip()

    # Claim number
    claim = _find(r"(?:claim\s+(?:number|no|ref)|reference\s+(?:number|no))\s*[:\-|]\s*(\S+)", text)
    if claim:
        fields["claim_number"] = claim.strip()

    # Manager / Supervisor
    manager = (_find(r"(?:manager|supervisor|reporting\s+officer)\s*[:\-|]\s*(.+)", text)
               or _find(r"(?:reported\s+to|person\s+reported\s+to)\s*[:\-|]\s*(.+)", text))
    if manager:
        fields["manager"] = manager.split("\n")[0].strip()

    # State — also try to infer from address
    state = _find(r"(?:state|state/territory)\s*[:\-|]\s*(VIC|NSW|QLD|TAS|SA|WA|ACT|NT)", text)
    if not state:
        # Try to infer from text mentions of state names
        state_map = {
            "victoria": "VIC", "new south wales": "NSW", "queensland": "QLD",
            "tasmania": "TAS", "south australia": "SA", "western australia": "WA",
        }
        text_lower = text.lower()
        for name_long, abbrev in state_map.items():
            if name_long in text_lower:
                state = abbrev
                break
    if state:
        fields["state"] = state.upper()

    return fields


# ---------------------------------------------------------------------------
# Public API
# ---------------------------------------------------------------------------

def parse_uploaded_report(file_bytes: bytes, file_type: str) -> dict:
    """
    Extract text from an uploaded file and parse fields.

    Parameters
    ----------
    file_bytes : raw bytes of the uploaded file
    file_type  : one of "pdf", "docx"

    Returns
    -------
    dict mapping wizard field names to extracted values
    """
    if file_type == "docx":
        text = _extract_text_from_docx(file_bytes)
    elif file_type == "pdf":
        text = _extract_text_from_pdf(file_bytes)
    else:
        return {}

    if not text.strip():
        return {}

    return _parse_fields_from_text(text)
